from django.shortcuts import render
from selenium import webdriver
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import pandas 
import time
import pypandoc
import io
from docx import Document
from bs4 import BeautifulSoup
import tempfile
import os
chrome_options = webdriver.ChromeOptions()
chrome_options.binary_location = os.environ.get("GOOGLE_CHROME_BIN")
chrome_options.add_argument("--headless")
chrome_options.add_argument("--disable-dev-shm-usage")
chrome_options.add_argument("--no-sandbox")

def auto_post(request):
    emailinput = request.POST["email"]
    passinput = request.POST["password"]
    platform = request.POST["service"]
    doc = request.FILES #returns a dict-like object
    file = doc['formFile']
    checker = str(file)
    print(checker)
    formFile = file.read()

    f_browser = webdriver.Chrome(executable_path=os.environ.get("CHROMEDRIVER_PATH"), chrome_options=chrome_options)

    if platform == "Wordpress":
        if checker.endswith('.html'):
            fd, path =tempfile.mkstemp(suffix = '.html')
            try:
                for chunk in file.chunks():
                    os.write(fd, chunk)
            except:
                raise Exception("Problem with the input file %s" % file.name)
            finally:
                os.close(fd)
            soup = BeautifulSoup(formFile, 'html.parser')
            try:
                titles = soup.title.string
            except:
                titles = "WordPress POST"

            try:
                for s in soup.select('title'):
                    s.extract()
            except:
                pass
        elif checker.endswith('.csv'):
            fd, path =tempfile.mkstemp(suffix = '.csv')
            try:
                for chunk in file.chunks():
                    os.write(fd, chunk)
            except:
                raise Exception("Problem with the input file %s" % file.name)
            finally:
                os.close(fd)
            parseFile = pandas.read_csv(path)
            test = parseFile.to_html()
            soup = BeautifulSoup(test, 'html.parser')
            try:
                titles = soup.title.string
            except:
                titles = "WordPress POST"

            try:
                for s in soup.select('title'):
                    s.extract()
            except:
                pass

        elif checker.endswith('.doc') or checker.endswith('.docx'):
            titledoc = Document(io.BytesIO(formFile))
            paraNr = 0
            for para in titledoc.paragraphs:
                if paraNr == 0:
                    titles=para.text 
                paraNr += 1
            print(titles)
            doc = pypandoc.convert_text(formFile, 'html',format='docx')
            soup = BeautifulSoup(doc, 'html.parser')
            try:
                for s in soup.select('title'):
                    s.extract()
            except:
                pass
        f_browser.get("https://stevenseo.com/adjacent/wp-admin/post-new.php")
        f_browser.maximize_window()
        time.sleep(5)
        element = WebDriverWait(f_browser, 15).until( 
        EC.presence_of_element_located((By.CSS_SELECTOR, '#user_login')) 
        )
        email=f_browser.find_element_by_id('user_login')
        email.send_keys(emailinput)
        password = f_browser.find_element_by_id('user_pass')
        password.send_keys(passinput+ Keys.RETURN)
        element = WebDriverWait(f_browser, 15).until( 
            EC.presence_of_element_located((By.XPATH, '//*[@id="title"]')) 
            )
        title = f_browser.find_element_by_xpath('//*[@id="title"]')
        title.send_keys(titles)
        button = f_browser.find_element_by_xpath('//*[@id="content-html"]')
        button.click()
        time.sleep(1)
        parag = f_browser.find_element_by_xpath('//*[@id="content"]')
        parag.send_keys(str(soup))
        time.sleep(2)
        button2 = f_browser.find_element_by_xpath('//*[@id="publish"]')
        testelem = f_browser.find_element_by_xpath('//*[@id="save-post"]')
        time.sleep(15)
        try:
            button2.click()
        except:
            time.sleep(2)
            webdriver.ActionChains(f_browser).move_to_element(testelem).perform()
            element = WebDriverWait(f_browser, 150).until( 
                EC.element_to_be_clickable((By.XPATH, '//*[@id="publish"]')) 
            )
            button2.click()
        time.sleep(10)
        f_browser.quit()
        return render(request, 'modal.html')

    elif platform == "Shopify":
        if checker.endswith('.html'):
            fd, path =tempfile.mkstemp(suffix = '.html')
            try:
                for chunk in file.chunks():
                    os.write(fd, chunk)
            except:
                raise Exception("Problem with the input file %s" % file.name)
            finally:
                os.close(fd)
            soup = BeautifulSoup(formFile, 'html.parser')
            try:
                titles = soup.title.string
            except:
                titles = "WordPress POST"

            try:
                for s in soup.select('title'):
                    s.extract()
            except:
                pass
        elif checker.endswith('.csv'):
            fd, path =tempfile.mkstemp(suffix = '.csv')
            try:
                for chunk in file.chunks():
                    os.write(fd, chunk)
            except:
                raise Exception("Problem with the input file %s" % file.name)
            finally:
                os.close(fd)
            parseFile = pandas.read_csv(path)
            test = parseFile.to_html()
            soup = BeautifulSoup(test, 'html.parser')
            try:
                titles = soup.title.string
            except:
                titles = "WordPress POST"

            try:
                for s in soup.select('title'):
                    s.extract()
            except:
                pass

        elif checker.endswith('.doc') or checker.endswith('.docx'):
            titledoc = Document(io.BytesIO(formFile))
            paraNr = 0
            for para in titledoc.paragraphs:
                if paraNr == 0:
                    titles=para.text 
                paraNr += 1

            doc = pypandoc.convert_text(formFile, 'html',format='docx')
            soup = BeautifulSoup(doc, 'html.parser')
            try:
                for s in soup.select('title'):
                    s.extract()
            except:
                pass
        f_browser.get("https://accounts.shopify.com/store-login")
        f_browser.maximize_window()
        element = WebDriverWait(f_browser, 15).until( 
            EC.presence_of_element_located((By.XPATH, '//*[@id="shop_domain"]')) 
            )
        store=f_browser.find_element_by_xpath('//*[@id="shop_domain"]')
        store.send_keys('align-right'+Keys.RETURN)
        time.sleep(3)
        element = WebDriverWait(f_browser, 15).until( 
            EC.presence_of_element_located((By.XPATH, '//*[@id="account_email"]')) 
            )
        email = f_browser.find_element_by_xpath('//*[@id="account_email"]')
        email.send_keys(emailinput+ Keys.RETURN)
        time.sleep(5)
        element = WebDriverWait(f_browser, 15).until( 
            EC.presence_of_element_located((By.XPATH, '//*[@id="account_password"]')) 
            )
        password = f_browser.find_element_by_xpath('//*[@id="account_password"]')
        password.send_keys(passinput+ Keys.RETURN)
        time.sleep(2)
        element = WebDriverWait(f_browser, 200).until( 
            EC.presence_of_element_located((By.CLASS_NAME, 'Polaris-Button__Text_yj3uv')) 
            )
        forward = f_browser.find_element_by_class_name('Polaris-Button__Text_yj3uv')
        forward.click()

        f_browser.get('https://align-right.myshopify.com/admin/products/new')

        element = WebDriverWait(f_browser, 200).until( 
            EC.presence_of_element_located((By.CLASS_NAME, 'Polaris-Button__Text_yj3uv')) 
            )
        forward = f_browser.find_element_by_class_name('Polaris-Button__Text_yj3uv')
        forward.click()

        element = WebDriverWait(f_browser, 15).until( 
            EC.presence_of_element_located((By.NAME, 'title')) 
            )

        title = f_browser.find_element_by_name('title')
        title.send_keys(titles)
        time.sleep(2)
        forward2 = f_browser.find_element_by_xpath('//button[@aria-label="Show HTML"]')
        forward2.click()
        element = WebDriverWait(f_browser, 15).until( 
            EC.presence_of_element_located((By.ID, 'product-description')) 
            )
        time.sleep(1)
        paragraph = f_browser.find_element_by_id('product-description')
        paragraph.send_keys(str(soup))
        forward3 = f_browser.find_element_by_xpath('//button[@aria-label="Save"]')
        try:
            forward3.click()
        except:
            element = WebDriverWait(f_browser, 200).until(
                EC.element_to_be_clickable((By.XPATH, '//button[@aria-label="Save"]')))
            forward3.click()
        time.sleep(10)
        f_browser.quit()
        return render(request, 'modal.html')

    elif platform == "Elementor":
        fd, path =tempfile.mkstemp(suffix = '.json')
        try:
            for chunk in file.chunks():
                os.write(fd, chunk)
        except:
            raise Exception("Problem with the input file %s" % file.name)
        finally:
            os.close(fd)

        f_browser.get('https://stevenseo.com/adjacent/wp-admin/edit.php?post_type=elementor_library&tabs_group=library')
        element = WebDriverWait(f_browser, 15).until( 
        EC.presence_of_element_located((By.CSS_SELECTOR, '#user_login')) 
        )
        email=f_browser.find_element_by_id('user_login')
        email.send_keys(emailinput)
        password = f_browser.find_element_by_id('user_pass')
        password.send_keys(passinput+ Keys.RETURN)
        element = WebDriverWait(f_browser, 200).until( 
            EC.presence_of_element_located((By.XPATH, '//*[@id="elementor-import-template-trigger"]')) 
            )
        time.sleep(1)
        button = f_browser.find_element_by_xpath('//*[@id="elementor-import-template-trigger"]')
        button.click()
        time.sleep(1)
        upload = f_browser.find_element_by_xpath('//input[@name="file"]')
        upload.send_keys(path)
        ele_button = f_browser.find_element_by_xpath('//input[@value="Import Now"]')
        ele_button.click()
        time.sleep(5)
        f_browser.quit()
        return render(request, 'modal.html')

    elif platform == "Divi":
        fd, path =tempfile.mkstemp(suffix = '.json')
        try:
            for chunk in file.chunks():
                os.write(fd, chunk)
        except:
            raise Exception("Problem with the input file %s" % file.name)
        finally:
            os.close(fd)
        f_browser.get('https://stevenseo.com/adjacent/wp-admin/edit.php?post_type=et_pb_layout')
        element = WebDriverWait(f_browser, 15).until( 
        EC.presence_of_element_located((By.CSS_SELECTOR, '#user_login')) 
        )
        email=f_browser.find_element_by_id('user_login')
        email.send_keys(emailinput)
        password = f_browser.find_element_by_id('user_pass')
        password.send_keys(passinput+ Keys.RETURN)
        time.sleep(1)
        button = f_browser.find_element_by_xpath('//a[@class="et-pb-portability-button"]')
        button.click()
        ele_button = f_browser.find_element_by_xpath('//a[@id="ui-id-2"]')
        ele_button.click()
        upload = f_browser.find_element_by_xpath('//input[@type="file"]')
        upload.send_keys(path)
        sender = f_browser.find_element_by_xpath('//a[@class="et-core-modal-action et-core-portability-import"]')
        sender.click()
        time.sleep(5)
        return render(request, 'modal.html')

def index(request):
    return render(request, 'index.html')